"""
MÓDULO: app/services/docx_generator.py - SERVIÇO DE MONTAGEM E FORMATAÇÃO DOCX

FUNÇÃO:
Define a classe `DocxGenerator`, responsável por transformar o JSON final
hierárquico (gerado pelo Agente 5) em um arquivo `.docx` formatado. O serviço
utiliza a biblioteca `python-docx` e lida com toda a estilização, a inserção
de metadados no cabeçalho e a aplicação correta da numeração hierárquica
(1., 2., 2.1., etc.).

ARQUITETURA:
A classe é um *Singleton* que armazena constantes de formatação (fontes,
cores, tamanhos) e fornece métodos utilitários (Helpers) para manipulação
avançada do DOCX (e.g., numeração de página, sombreamento de célula).

RESPONSABILIDADES CHAVE:
1. **Geração do Cabeçalho:** Cria a tabela fixa do cabeçalho do documento,
   inserindo logo, codificação, revisão e o campo dinâmico de numeração
   de página (Página X de Y).
2. **Numeração Automática:** Percorre a estrutura `corpo_documento` do JSON,
   aplicando numeração sequencial correta (Seção principal `secao_counter`,
   subseção `subsecao_counter`).
3. **Estilização:** Aplica formatação consistente (fonte Arial, tamanhos
   específicos, recuos) aos títulos e ao conteúdo do corpo.
4. **Persistência:** Salva o arquivo `.docx` no diretório de saídas
   configurado (`settings.OUTPUTS_PATH`).

FLUXO DE USO:
O `ChatOrchestrator` chama a função `create_document` na última etapa
do fluxo, passando o objeto `DocumentoFinalJSON` completo.
"""
import logging
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importa nossas configurações e o schema HIERÁRQUICO
from app.core.config import settings
from app.core.schemas import DocumentoFinalJSON

logger = logging.getLogger(__name__)

class DocxGenerator:
    """
    Serviço "Montador" (o último passo do fluxo).
    Responsável por pegar o DocumentoFinalJSON e gerar o arquivo .docx final.
    """
    
    def __init__(self):
        # Constantes de formatação, definidas uma única vez
        self.FONT_NAME = 'Arial'
        self.FONT_SIZE_LABEL = Pt(10)
        self.FONT_SIZE_VALUE = Pt(12)
        # Largura total da tabela do cabeçalho
        self.TABLE_TOTAL_WIDTH = Cm(15.92)
        self.COR_LARANJA_CABECALHO = "F79646"
        self.COR_PRETA_TEXTO = RGBColor(0x00, 0x00, 0x00)

    # --- INÍCIO DAS FUNÇÕES HELPER (Privadas) ---
    
    def _add_page_number_field(self, paragraph, bold=False, font_name='Arial', font_size=Pt(12)):
        """
        Adiciona o campo de numeração de página dinâmico (Página X de Y)
        ao parágrafo, usando o XML de baixo nível do DOCX.
        Isso garante que a numeração seja atualizada automaticamente pelo Word.
        """
        paragraph.clear() 
        # Campo para o número da página atual (PAGE)
        fldChar_begin_page = OxmlElement('w:fldChar')
        fldChar_begin_page.set(qn('w:fldCharType'), 'begin')
        instrText_page = OxmlElement('w:instrText')
        instrText_page.set(qn('xml:space'), 'preserve')
        instrText_page.text = 'PAGE \\* MERGEFORMAT'
        fldChar_separate_page = OxmlElement('w:fldChar')
        fldChar_separate_page.set(qn('w:fldCharType'), 'separate')
        fldChar_end_page = OxmlElement('w:fldChar')
        fldChar_end_page.set(qn('w:fldCharType'), 'end')
        run_page = paragraph.add_run()
        run_page._r.append(fldChar_begin_page)
        run_page._r.append(instrText_page)
        run_page._r.append(fldChar_separate_page)
        run_page._r.append(fldChar_end_page)
        run_page.font.name = font_name
        run_page.font.size = font_size
        run_page.font.bold = bold
        
        # O texto " de "
        run_de = paragraph.add_run(' de ')
        run_de.font.name = font_name
        run_de.font.size = font_size
        run_de.font.bold = bold
        
        # Campo para o número total de páginas (NUMPAGES)
        fldChar_begin_num = OxmlElement('w:fldChar')
        fldChar_begin_num.set(qn('w:fldCharType'), 'begin')
        instrText_numpages = OxmlElement('w:instrText')
        instrText_numpages.set(qn('xml:space'), 'preserve')
        instrText_numpages.text = 'NUMPAGES \\* MERGEFORMAT'
        fldChar_end_num = OxmlElement('w:fldChar')
        fldChar_end_num.set(qn('w:fldCharType'), 'end')
        run_numpages = paragraph.add_run()
        run_numpages._r.append(fldChar_begin_num)
        run_numpages._r.append(instrText_numpages)
        run_numpages._r.append(fldChar_end_num)
        run_numpages.font.name = font_name
        run_numpages.font.size = font_size
        run_numpages.font.bold = bold

    def _set_cell_text(self, cell, text, bold=False, size=12, font_name='Arial', align='LEFT'):
        """Define o texto, a formatação e o alinhamento de uma célula."""
        cell.text = text
        paragraph = cell.paragraphs[0]
        
        # Define o alinhamento do parágrafo
        if align.upper() == 'CENTER':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align.upper() == 'RIGHT':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # Define a formatação do Run (texto)
        run = paragraph.runs[0]
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        return paragraph

    def _set_cell_nowrap(self, cell):
        """Impede que o conteúdo de uma célula quebre linha (nowrap)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        noWrap = OxmlElement('w:noWrap')
        tcPr.append(noWrap)

    def _set_cell_label_value(self, cell, label, value, bold_label=False, bold_value=False, align='CENTER'):
        """
        Formata uma célula com duas linhas: um Label (título) e um Value (valor).
        Usado para Codificação, Data de Revisão, Revisão e Página no cabeçalho.
        """
        # Parágrafo para o Label
        p_label = cell.paragraphs[0]
        p_label.text = label
        if align.upper() == 'CENTER':
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.runs[0]
        run_label.font.name = self.FONT_NAME
        run_label.font.size = self.FONT_SIZE_LABEL
        run_label.font.bold = bold_label
        
        # Parágrafo para o Value (adicionado como um novo parágrafo)
        p_value = cell.add_paragraph() 
        if align.upper() == 'CENTER':
            p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Se o label for "Página", insere o campo dinâmico de numeração
        if label == 'Página':
            self._add_page_number_field(p_value, bold=bold_value, font_name=self.FONT_NAME, font_size=self.FONT_SIZE_VALUE)
        else:
            # Insere o valor como texto normal
            run_value = p_value.add_run(value) 
            run_value.font.name = self.FONT_NAME
            run_value.font.size = self.FONT_SIZE_VALUE
            run_value.font.bold = bold_value
            
        # Remove parágrafos excedentes (limita a 2: label e value)
        while len(cell.paragraphs) > 2:
            p = cell.paragraphs[-1]
            cell._element.remove(p._p)
            
        return p_label, p_value 

    def _add_image_to_cell(self, cell, image_path, width_cm=None):
        """Adiciona uma imagem centralizada em uma célula, com tratamento de erro."""
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = paragraph.add_run()
        try:
            # Converte o objeto Path para string antes de usar no add_picture
            if width_cm:
                run.add_picture(str(image_path), width=Cm(width_cm)) 
            else:
                run.add_picture(str(image_path)) 
        except FileNotFoundError:
            logger.warning(f"Arquivo de imagem não encontrado em: {image_path}. Inserindo placeholder.")
            run.text = "[LOGO]"
        except Exception as e:
            logger.error(f"Erro ao inserir imagem {image_path}: {e}")
            run.text = "[ERRO IMAGEM]"


    def _set_cell_shading(self, cell, hex_color_str):
        """Aplica sombreamento de fundo (background) a uma célula."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color_str)
        tcPr.append(shd)

    def _add_historico_revisoes(self, document):
        """
        Adiciona a primeira seção fixa do documento: "1. Histórico das Revisões"
        e insere o template da tabela de revisões.
        """
        # Título "1. Histórico das Revisões"
        p_heading = document.add_paragraph()
        run = p_heading.add_run('1. Histórico das Revisões') # Número 1 Fixo
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE_VALUE
        run.font.bold = True 
        run.font.color.rgb = self.COR_PRETA_TEXTO
        
        p_format = p_heading.paragraph_format
        p_format.left_indent = Cm(0.5)
        p_format.space_before = Pt(12) 
        p_format.space_after = Pt(6) 
        
        # Cria a tabela de revisões
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.width = self.TABLE_TOTAL_WIDTH 
        
        hdr_cells = table.rows[0].cells
        titulos = ['Revisão', 'Data', 'Alteração', 'Revisor', 'Validação']
        
        # Configura o cabeçalho da tabela com sombreamento laranja
        for i, titulo in enumerate(titulos):
            cell = hdr_cells[i]
            p = self._set_cell_text(cell, titulo, bold=True, align='CENTER')
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # Texto Branco
            self._set_cell_shading(cell, self.COR_LARANJA_CABECALHO)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Adiciona uma linha vazia para preenchimento futuro (template)
        table.add_row()
        
        # Define as larguras fixas das colunas
        table.columns[0].width = Cm(1.5)
        table.columns[1].width = Cm(2.5)
        table.columns[2].width = Cm(6.42)
        table.columns[3].width = Cm(2.75)
        table.columns[4].width = Cm(2.75)

    # --- FIM DAS FUNÇÕES HELPER ---

    # --- FUNÇÃO PRINCIPAL DO SERVIÇO ---
    def create_document(self, data: DocumentoFinalJSON) -> str:
        """
        Gera o documento DOCX completo a partir dos dados do JSON.

        Args:
            data: Objeto DocumentoFinalJSON contendo metadados e corpo hierárquico.

        Returns:
            O caminho completo do arquivo salvo (string).
        """
        logger.info(f"Iniciando a criação do documento: {data.codificacao}.docx")
        
        document = Document()

        # 1. Configuração de Estilo e Margens do Documento
        style = document.styles['Normal']
        font = style.font
        font.name = self.FONT_NAME
        font.size = self.FONT_SIZE_VALUE
        
        section = document.sections[0]
        # Margens A4 padrão
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.header_distance = Pt(0) # Remove espaço extra do cabeçalho
        
        # 2. Construção do Cabeçalho
        header = section.header
        if header.paragraphs:
            header_paragraph = header.paragraphs[0]
            header_format = header_paragraph.paragraph_format
            header_format.space_before = Pt(0)
            header_format.space_after = Pt(0)
        
        header_table = header.add_table(rows=3, cols=5, width=self.TABLE_TOTAL_WIDTH)
        header_table.style = 'Table Grid'
        
        # Define as larguras das colunas do cabeçalho
        header_table.columns[0].width = Cm(3.5)
        header_table.columns[1].width = Cm(4.71)
        header_table.columns[2].width = Cm(4.71)
        header_table.columns[3].width = Cm(1.5)
        header_table.columns[4].width = Cm(1.5)
        
        # Insere o Logo na célula (0, 0)
        logo_full_path = settings.ASSETS_PATH / data.logo_path
        self._add_image_to_cell(header_table.cell(0, 0), logo_full_path, width_cm=3.0)
        
        # Tipo de Documento (célula 0, 1)
        p = self._set_cell_text(header_table.cell(0, 1), data.tipo_documento, 
                             bold=False, size=12, align='CENTER') 
        header_table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Codificação (célula 1, 1)
        cell_cod = header_table.cell(1, 1)
        self._set_cell_label_value(cell_cod, 'Codificação', data.codificacao, 
                             bold_label=False, bold_value=True)
        self._set_cell_nowrap(cell_cod)
        
        # Data de Revisão (célula 1, 2)
        cell_data = header_table.cell(1, 2)
        self._set_cell_label_value(cell_data, 'Data de Revisão', data.data_revisao,
                             bold_label=False, bold_value=True)
        self._set_cell_nowrap(cell_data)
                             
        # Número de Revisão (célula 1, 3)
        self._set_cell_label_value(header_table.cell(1, 3), 'Revisão', data.numero_revisao,
                             bold_label=False, bold_value=True)
        
        # Numeração de Página (célula 1, 4) - Usa a função com campo dinâmico
        p_label, p_value = self._set_cell_label_value(header_table.cell(1, 4), 'Página', '',
                                               bold_label=False, bold_value=True)

        # Título do Documento (célula 2, 1)
        p_title = self._set_cell_text(header_table.cell(2, 1), data.titulo_documento, 
                                 bold=True, size=12, align='CENTER')
        header_table.cell(2, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Merge de células no cabeçalho (Lógica de layout complexa)
        header_table.cell(0, 0).merge(header_table.cell(1, 0)).merge(header_table.cell(2, 0))
        header_table.cell(0, 1).merge(header_table.cell(0, 2)).merge(header_table.cell(0, 3)).merge(header_table.cell(0, 4))
        header_table.cell(2, 1).merge(header_table.cell(2, 2)).merge(header_table.cell(2, 3)).merge(header_table.cell(2, 4))

        # --- 3. Corpo do Documento ---
        
        # Adiciona a primeira seção fixa: "1. Histórico das Revisões"
        self._add_historico_revisoes(document) 
        
        document.add_page_break() # Quebra de página após o histórico
        
        # Espaçador para garantir que o texto não fique colado no cabeçalho após o Page Break
        spacer_p = document.add_paragraph()
        spacer_format = spacer_p.paragraph_format
        spacer_format.space_before = Pt(0)
        spacer_format.space_after = Pt(0)
        spacer_run = spacer_p.add_run()
        spacer_run.font.size = Pt(12) 
        
        # Início da numeração hierárquica a partir do número 2
        secao_counter = 2 
        is_first_section_after_break = True
        
        # Loop principal para iterar sobre as Seções
        for secao in data.corpo_documento:
            
            # --- Título da Seção (Nível 1) ---
            # O Python adiciona o número (ex: "2. Objetivo")
            titulo_secao = f"{secao_counter}. {secao.titulo}"
            p_heading = document.add_paragraph()
            run_heading = p_heading.add_run(titulo_secao)
            run_heading.font.name = self.FONT_NAME
            run_heading.font.size = self.FONT_SIZE_VALUE
            run_heading.font.bold = True 
            run_heading.font.color.rgb = self.COR_PRETA_TEXTO
            
            p_format = p_heading.paragraph_format
            p_format.left_indent = Cm(0.5) # Recuo para a Seção Nível 1
            
            # Controla o espaçamento antes da primeira seção após o Page Break
            if is_first_section_after_break:
                p_format.space_before = Pt(0)
                is_first_section_after_break = False
            else:
                p_format.space_before = Pt(12)
            p_format.space_after = Pt(6)

            # --- Conteúdo da Seção (Texto Principal) ---
            p_text = document.add_paragraph(secao.conteudo)
            if not secao.conteudo: p_text.text = ""
            # Adiciona formatação ao texto
            run_text = p_text.runs[0] if p_text.runs else p_text.add_run("")
            run_text.font.name = self.FONT_NAME
            run_text.font.size = self.FONT_SIZE_VALUE
            p_text.paragraph_format.left_indent = Cm(0.5) # Recua o texto principal (mesmo recuo do título)

            # --- Loop de Subseções (Nível 2) ---
            subsecao_counter = 1
            for subsecao in secao.subsecoes:
                # O Python adiciona o número (ex: "2.1. Fluxograma")
                titulo_subsecao = f"{secao_counter}.{subsecao_counter}. {subsecao.titulo}"
                p_sub_heading = document.add_paragraph()
                run_sub = p_sub_heading.add_run(titulo_subsecao)
                run_sub.font.name = self.FONT_NAME
                run_sub.font.size = self.FONT_SIZE_VALUE
                run_sub.font.bold = True 
                run_sub.font.color.rgb = self.COR_PRETA_TEXTO
                
                p_sub_format = p_sub_heading.paragraph_format
                p_sub_format.left_indent = Cm(1.0) # Recuo maior para Nível 2
                p_sub_format.space_before = Pt(10)
                p_sub_format.space_after = Pt(4)
                
                # Conteúdo da Subseção (Texto)
                p_sub_text = document.add_paragraph(subsecao.conteudo)
                if not subsecao.conteudo: p_sub_text.text = ""
                # Adiciona formatação ao texto da subseção
                run_sub_text = p_sub_text.runs[0] if p_sub_text.runs else p_sub_text.add_run("")
                run_sub_text.font.name = self.FONT_NAME
                run_sub_text.font.size = self.FONT_SIZE_VALUE
                p_sub_text.paragraph_format.left_indent = Cm(1.0) # Recua o texto (mesmo recuo do título)
                
                subsecao_counter += 1 # Incrementa o contador da subseção
            
            secao_counter += 1 # Incrementa o contador da seção principal
        
        # --- 4. Salvar Documento ---
        # Garante que o diretório de saída exista
        settings.OUTPUTS_PATH.mkdir(parents=True, exist_ok=True)
        
        # Define o nome e o caminho completo do arquivo
        output_filename = f"{data.codificacao}.docx"
        output_path = settings.OUTPUTS_PATH / output_filename
        
        # Salva o documento no disco
        document.save(output_path)
        logger.info(f"Documento '{output_path}' criado com sucesso!")
        
        # Retorna o caminho como string para o Orquestrador
        return str(output_path)


# Cria uma instância única do serviço (Singleton)
docx_service = DocxGenerator()